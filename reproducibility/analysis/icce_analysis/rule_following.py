from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from scipy import stats

from .config import FEEDBACK_FILENAMES, GROUP_LABELS, RULE_LABELS
from .io_utils import load_reclassification, normalize_email, normalize_group
from .stats_utils import cohens_d


RULE_KEYS = tuple(RULE_LABELS)


def _yes_no(value) -> int | None:
    s = str(value or "").strip().lower()
    if s in {"yes", "y", "1", "true", "بله", "آره", "اره"} or s.startswith("yes"):
        return 1
    if s in {"no", "n", "0", "false", "خیر", "نه"} or s.startswith("no"):
        return 0
    return None


def _detect_columns(headers: list[str]) -> dict[str, int | None]:
    h = [x.strip().lower() for x in headers]

    def find(words):
        for i, text in enumerate(h):
            if any(word in text for word in words):
                return i
        return None

    return {
        "student": find(["student", "email"]),
        "summary": find(["summary"]),
        "R1": find(["rule 1", "prioritize", "process"]),
        "R2": find(["rule 2", "concept"]),
        "R3": find(["rule 3", "hint", "step"]),
        "R4": find(["rule 4", "active"]),
        "R5": find(["rule 5", "critical", "evaluation"]),
        "R6": find(["rule 6", "ethical", "no-help"]),
    }


def parse_feedback_docx(path: Path) -> list[dict]:
    match = re.search(r"feedback_group([23])_practice([234])", path.stem.lower())
    if not match:
        raise ValueError(f"Cannot infer assigned group/practice from {path.name}")

    assigned_group = f"Group {match.group(1)}"
    practice = int(match.group(2))

    doc = Document(path)
    if not doc.tables:
        return []

    table = doc.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    cols = _detect_columns(headers)

    records = []
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        idx = cols["student"]
        email = normalize_email(cells[idx] if idx is not None and idx < len(cells) else "")
        if not email:
            continue

        rec = {
            "email": email.replace("_gmail.com", "@gmail.com").replace("_ut.ac.ir", "@ut.ac.ir"),
            "practice": practice,
            "assigned_group": assigned_group,
            "summary": cells[cols["summary"]] if cols["summary"] is not None and cols["summary"] < len(cells) else "",
        }
        for rule in RULE_KEYS:
            col = cols[rule]
            rec[rule] = _yes_no(cells[col]) if col is not None and col < len(cells) else None
        rec["total_yes"] = sum(v for v in (rec[r] for r in RULE_KEYS) if v is not None)
        records.append(rec)
    return records


def _student_summary(records: pd.DataFrame, group_col: str) -> pd.DataFrame:
    rows = []
    for email, d in records.groupby("email"):
        group = d[group_col].iloc[0]
        row = {
            "email": email,
            "group": group,
            "n_practices_coded": len(d),
            "mean_total_yes": d["total_yes"].mean(),
            "sd_total_yes": d["total_yes"].std(ddof=1),
        }
        for rule in RULE_KEYS:
            row[rule] = d[rule].mean()
        rows.append(row)
    return pd.DataFrame(rows)


def _overall_comparison(student_df: pd.DataFrame) -> dict:
    unrestricted = student_df.loc[student_df["group"] == "Group 2", "mean_total_yes"].dropna().to_numpy(float)
    guided = student_df.loc[student_df["group"] == "Group 3", "mean_total_yes"].dropna().to_numpy(float)
    t = stats.ttest_ind(guided, unrestricted, equal_var=False)
    return {
        "unrestricted_n": len(unrestricted),
        "guided_n": len(guided),
        "unrestricted_mean": unrestricted.mean(),
        "guided_mean": guided.mean(),
        "difference_guided_minus_unrestricted": guided.mean() - unrestricted.mean(),
        "cohens_d": cohens_d(unrestricted, guided),
        "welch_p": float(t.pvalue),
    }


def analyze_rule_following(data_dir: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    feedback_dir = data_dir / "feedback"
    records = []
    for filename in FEEDBACK_FILENAMES:
        path = feedback_dir / filename
        if not path.exists():
            raise FileNotFoundError(f"Missing feedback file: {path}")
        records.extend(parse_feedback_docx(path))

    df = pd.DataFrame(records)
    reclass = load_reclassification(data_dir / "reclassification.csv")
    df["enacted_group"] = [
        reclass.get(email, assigned)
        for email, assigned in zip(df["email"], df["assigned_group"])
    ]
    df.to_csv(output_dir / "rule_following_records_long.csv", index=False)

    assigned_students = _student_summary(df, "assigned_group")
    enacted_students = _student_summary(df, "enacted_group")
    assigned_students.to_csv(output_dir / "student_compliance_assigned.csv", index=False)
    enacted_students.to_csv(output_dir / "student_compliance_enacted.csv", index=False)

    comparisons = pd.DataFrame([
        {"grouping": "assigned", **_overall_comparison(assigned_students)},
        {"grouping": "enacted", **_overall_comparison(enacted_students)},
    ])
    comparisons.to_csv(output_dir / "overall_compliance_comparison.csv", index=False)

    # Rule-level percentages use transcript records and assigned condition.
    rule_rows = []
    for rule in RULE_KEYS:
        row = {"rule": rule, "rule_label": RULE_LABELS[rule]}
        table = []
        for group in ["Group 2", "Group 3"]:
            values = df.loc[df["assigned_group"] == group, rule].dropna().astype(int)
            yes = int(values.sum())
            n = int(len(values))
            row[f"{GROUP_LABELS[group]}_yes"] = yes
            row[f"{GROUP_LABELS[group]}_n"] = n
            row[f"{GROUP_LABELS[group]}_rate"] = yes / n if n else np.nan
            table.append([yes, n - yes])
        _, p = stats.fisher_exact(table, alternative="two-sided")
        row["fisher_p"] = float(p)
        rule_rows.append(row)
    pd.DataFrame(rule_rows).to_csv(output_dir / "rule_level_compliance_assigned.csv", index=False)

    # Practice-specific mean rule score.
    practice_rows = []
    for practice in [2, 3, 4]:
        d = df[df["practice"] == practice]
        a = d.loc[d["assigned_group"] == "Group 2", "total_yes"].to_numpy(float)
        b = d.loc[d["assigned_group"] == "Group 3", "total_yes"].to_numpy(float)
        test = stats.ttest_ind(b, a, equal_var=False)
        practice_rows.append({
            "practice": practice,
            "unrestricted_n": len(a),
            "guided_n": len(b),
            "unrestricted_mean": a.mean() if len(a) else np.nan,
            "guided_mean": b.mean() if len(b) else np.nan,
            "difference": b.mean() - a.mean() if len(a) and len(b) else np.nan,
            "welch_p": float(test.pvalue) if len(a) > 1 and len(b) > 1 else np.nan,
        })
    pd.DataFrame(practice_rows).to_csv(output_dir / "practice_compliance_assigned.csv", index=False)

    print(f"[rules] wrote outputs to {output_dir}")
